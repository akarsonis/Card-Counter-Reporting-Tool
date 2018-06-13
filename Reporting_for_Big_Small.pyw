import random

import time
from datetime import datetime, timedelta
import os
import sys

from tkinter import *

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.shared import Inches

import xlrd
import openpyxl

from collections import Counter
import re

from bs4 import BeautifulSoup
import requests

#Goose function

def rapping_goose():
    goose_something = Toplevel()
    goose_window = Canvas(goose_something, width = 460, height = 380)
    goose_window.pack(expand = NO, fill = BOTH)
    goose_pic = PhotoImage(file = 'JWakw7w.gif')

    goose_window.create_image(50, 10, image = goose_pic, anchor = NW)
    goose_window.gif1 = goose_pic

intro = 'Ezugi Fraud and Risk department detection tools, reports and operations are based on a search for particular betting patterns to identify advantage play or cheating behavior within Ezugi gaming network. The account that is a subject of this report is under Ezugi Fraud and Risk department\'s observation since the player\'s gameplay revealed signs of Card Counting.'

conclusion = 'Ezugi Fraud and Risk department\'s advice is to block the player from Live Casino games since \"UID\" is not joining Live Casino to gamble, but rather is trying to find an advantage over the house.'

p2a = 'First, during the investigated time frame the player participated in game_count games of game_names and did not show any preference to a specific dealer.'
p2b = 'To begin, during the investigated time frame the player \"UID\" participated in games of game_names. The player held game_count games in total and did not show any preference to a specific dealer.'
p2c = 'To start, during the investigated period the player held games of game_names and in total participated in game_count games. Player did not show a preference to any specific dealer.'


p3a = 'When it comes to the betting pattern, in majority of games the user preferred placing small wagers on main betting spots which generally were small_wager, however once it was the end of the shoe \"UID\'s\" wager could increase greatly to up to big_wager. Moreover, during these situations the wagers were placed on \"Big\" or \"Small\" side bets only.'

p3b = 'When looking at the player\'s betting pattern, in majority of games the user preferred placing small wagers on main betting spots which generally were small_wager, however once it was the end of the shoe \"UID\'s\" wager could increase greatly to up to big_wager. Moreover, the high wagers were placed on \"Big\" or \"Small\" side bets only.'

p3c = 'As regards to the betting strategy, in majority of games the user preferred placing small wagers on main betting spots which generally were small_wager, however once it was the end of the shoe \"UID\'s\" wager was significantly increased to up to big_wager. Moreover, during these situations the wagers were placed on \"Big\" or \"Small\" side bets only.'


p4a = 'To continue, due to the players chaotic betting and high bet ramps, an investigation was performed in order to reveal if the user was involved in any kind of illicit activity. Investigation results showed that, the player was counting cards such as 8s and 9s. Please see examples further:'

p51a = 'The key is that for the player to win the Big bet, the game has to end with the 5 or 6 cards on the table. So, the more high cards are out the higher the probability for the game to end up with 5 or 6 cards on the table and bring the winning outcome for the player and vice-versa for \"Small\" bet.'

p52a = 'The key is that for the player to win the \"Small\" bet, the game has to end with the 4 cards on the table. So, the richer the deck is with the high cards the higher the probability for the game to end up with 4 cards on the table and bring the winning outcome for the player and vice-versa for \"Big\" bet.'

workbook = xlrd.open_workbook('Example.xlsx')

worksheet = workbook.sheet_by_index(0)

document = Document('Ezugi_template_Big_Small.docx')

root = Tk()
root.title('Player Reporting')
root.geometry('650x310')
root.iconbitmap('fish.ico')

table0 = document.tables[0]   # info about the player

cell011 = table0.cell(1,1)    # operator ID

cell021 = table0.cell(2,1)    # screen name

cell031 = table0.cell(3,1)    # UID

cell040 = table0.cell(4,0)    # agg. win/loss

cell041 = table0.cell(4,1)    # profit

cell051 = table0.cell(5,1)    # turnover

cell061 = table0.cell(6,1)    # Margin

cell071 = table0.cell(7,1)    # date of enrollment

cell081 = table0.cell(8,1)
#
table1 = document.tables[1]   # Body of a text

cell110 = table1.cell(1,0)
#
table2 = document.tables[2]   # Conclusion

cell210 = table2.cell(1,0)
#

run0 = document.paragraphs[13].add_run(os.environ['USERNAME'])
font0 = run0.font
font0.size = Pt(15)


if os.environ['USERNAME'] == 'Aleksandrs':
    run01 = document.paragraphs[13].add_run(' Karsonis')
    font01 = run01.font
    font01.size = Pt(14)    
    if os.environ['USERNAME'] == 'Arturs':
        run02 = document.paragraphs[13].add_run(' Lusis')
        font02 = run01.font
        font02.size = Pt(14)
        if os.environ['USERNAME'] == 'Alina':
            run02 = document.paragraphs[13].add_run(' Heifeca')
            font02 = run01.font
            font02.size = Pt(14)        

run2 = document.paragraphs[15].add_run(time.strftime("%d.%m.%y"))
font2 = run2.font
font2.size = Pt(14)

#Max bet function xl

def follow_up():
   col = 12
   return max(worksheet.col_values(col, start_rowx=4))

big_wager = follow_up()
big_wager = "{:6,.2f}".format(big_wager)

#Turnover function xl

def turnover():
    col = 12
    return sum(worksheet.col_values(col, start_rowx=4))

turnover_xl = turnover()

#Payoff function xl

def payoff():
    col = 13
    return sum(worksheet.col_values(col, start_rowx=4))

payoff_xl = payoff()

#Function MAIN

def Ezugi_function(event):
    
    global p2a
    global p2b
    global p2c
    global p3a
    global p3b
    global p3c
    global conclusion
    global big_wager
    
    if dates_entry.get() == '':
        run3 = document.paragraphs[14].add_run('All history')
        font2 = run3.font
        font2.size = Pt(14)
    else:
        run3 = document.paragraphs[14].add_run(dates_entry.get())
        font2 = run3.font
        font2.size = Pt(14)
    
    # BIG first login function ##############################################
    
    date_count = 0
    for row in range(worksheet.nrows):
        if (worksheet.cell_value(row, 10)) != '':
            date_count = date_count + 1
            
    date_count = date_count + 2
    
    test_date = worksheet.cell_value(date_count, 10)
    
    enrolled = xlrd.xldate_as_tuple(test_date, workbook.datemode)
    
    enrolled = list(enrolled)
    
    del enrolled[3]
    del enrolled[3]
    del enrolled[3]
    
    enrolled[0], enrolled[1], enrolled[2] = enrolled[2], enrolled[1], enrolled[0]
    
    enrolled = str(enrolled)
    enrolled = enrolled.replace('[', '')
    enrolled = enrolled.replace(']', '')
    enrolled = enrolled.replace(', ', '.')
    
    if enrolled[0:2].find('.') == 1:
        enrolled = '0' + enrolled
    
    if enrolled[3:5].find('.') == 1:
        enrolled = enrolled[0:3] + '0' + enrolled[3:]
    
    cell071.text = str(enrolled)    
    
    #GETTING SCREEN NAME AND BRAND WITH BEAUTIFUL SOUP AND XL
    
    #Logging in
    
    with requests.Session() as c:
        url = 'https://sbo.ezugi.com/office.php?page=login'
        USERNAME = 'akarsonis'
        PASSWORD = 'Ezugi123456'
        c.get(url)
        login_data = dict(username=USERNAME, password=PASSWORD, 
                          language_view='english', submit='Login')
        c.post(url, data=login_data, headers={'Referer': 'https://sbo.ezugi.com/office.php?page=login'})
        page = c.get(link_entry.get())
    
    #Variables
    
    soup = BeautifulSoup(page.content, 'html.parser')
    page_itself = soup.prettify()
    all_imgs = soup.find_all('img', title=True)[10:12]
    all_titles = [one_title['title'] for one_title in all_imgs]
    
    #text manipulation
    
    if len(all_titles[0]) > len(all_titles[1]):
        del all_titles[1]
    else:
        del all_titles[0]
    
    all_titles = str(all_titles)
    anchor = " - "
    all_titles = all_titles.split(anchor)[1].split(anchor)[0]
    anchor_operator = "Operator: "
    global operatorid
    operatorid = all_titles.split(anchor_operator, 1)[-1]
    global screen_name
    screen_name = all_titles.split(anchor_operator, 1)[0]
    screen_name = screen_name[:-3]  
    
    #inserting Screen name and operator ID into doc
    
    global operator_id_digits
    operator_id_digits = worksheet.cell(6,3).value
    operator_id_digits = int(operator_id_digits)
    
    cell021.text = screen_name
    cell011.text = operatorid + ' ' + str(operator_id_digits)
      
    #UID from excel
    global uid
    uid = worksheet.cell(6,6).value          
    
    if isinstance(uid, str):
      cell031.text = uid
    else:
      uid = int(uid)
      uid = str(uid)
      cell031.text = uid

    #Margin and currency
    
    user_currency = str(worksheet.cell(6,14).value)
    user_currency = ' ' + user_currency
    
    user_turnover = turnover_xl
    user_turnover = float(user_turnover)
    cell051.text = "{:6,.2f}".format(user_turnover) + user_currency    
   
    user_profit = float(payoff_xl) - float(turnover_xl)
    user_profit = "{:6,.2f}".format(user_profit)
    user_profit = str(user_profit)
    user_profit = user_profit.replace(' ', '')
    cell041.text = user_profit + user_currency
    user_profit = user_profit.replace(',', '')
    
    if float(user_profit) > 0 or float(user_profit) == 0:
        cell040.paragraphs[0].add_run('Aggregated win:').bold = True
    else:
        cell040.paragraphs[0].add_run('Aggregated loss:').bold = True
    
    user_profit = float(user_profit)
    user_turnover = float(user_turnover)
    margin = (user_profit / user_turnover) * 100
    margin = "{:6,.2f}".format(margin)
    margin = str(margin)
    margin = margin.replace(' ', '')      
    cell061.text = margin + ' %'
    
    #1st and 2nd paragraphs
    
    intro_paragraph = cell110.paragraphs[0]
    intro_paragraph.text = intro
    
    # Big first paragraph function
    
    all_game_types = []
    
    bac_count = 0
    for row in range(worksheet.nrows):
        if 'accara' in (worksheet.cell_value(row, 7)):
            bac_count = bac_count + 1
    
    if bac_count > 0:
        all_game_types.append('Baccarat')
    
    rul_count = 0
    for row in range(worksheet.nrows):
        if 'oule' in (worksheet.cell_value(row, 7)):
            rul_count = rul_count + 1
    
    if rul_count > 0:
        all_game_types.append('Roulette')
    
    bla_count = 0
    for row in range(worksheet.nrows):
        if 'jack' in (worksheet.cell_value(row, 7)):
            bla_count = bla_count + 1
    
    if bla_count > 0:
        all_game_types.append('Blackjack')
    
    hol_count = 0
    for row in range(worksheet.nrows):
        if 'hold' in (worksheet.cell_value(row, 7)):
            hol_count = hol_count + 1
    
    if hol_count > 0:
        all_game_types.append('Casino Holdem')
    
    if len(all_game_types) > 1:
        all_game_types.insert(-1, 'and')

    all_game_types = ', '.join(all_game_types[:-2]) +' '+ ' '.join(all_game_types[-2:])
    all_game_types = str(all_game_types)

    all_game_types = all_game_types.replace('[', '')
    all_game_types = all_game_types.replace(']', '')
    all_game_types = all_game_types.replace('\'', '')
    
    if len(all_game_types) < 12:
        all_game_types = all_game_types.replace(' ', '')
        all_game_types = all_game_types + ' only'

    p2a = p2a.replace('game_names', all_game_types)
    p2b = p2b.replace('game_names', all_game_types)
    p2c = p2c.replace('game_names', all_game_types)
    
    #game count in 2nd paragraph
    
    game_count = 0
    for row in range(worksheet.nrows):
        if (worksheet.cell_value(row, 11)) == 'Table Bet':
            game_count = game_count + 1
    
    game_count = game_count
    game_count = "{:,}".format(game_count)
    game_count = str(game_count)
    game_count = game_count.replace(' ', '')  
    
    p2a = p2a.replace('game_count', game_count)
    p2b = p2b.replace('game_count', game_count)
    p2c = p2c.replace('game_count', game_count)
    
    p2b = p2b.replace('UID', screen_name)
    
    p2_random = random.choice([p2a, p2b, p2c])
    p2_paragraph = cell110.paragraphs[1]
    p2_paragraph.text = p2_random    

    #p3
    #Calculating in xl and inserting in doc smallest wager placed
    
    player_bets = []
    
    for row in range(worksheet.nrows):
        player_bets.append(worksheet.cell_value(row, 12))
    
    del player_bets[0:4]
    
    while 0.0 in player_bets:
        player_bets.remove(0.0)
    
    most_often_small_wager = Counter(player_bets).most_common(1)
    
    most_often_small_wager = str(most_often_small_wager)
    
    most_often_small_wager = most_often_small_wager.replace('[', '')
    most_often_small_wager = most_often_small_wager.replace(']', '')
    most_often_small_wager = most_often_small_wager.replace('(', '')
    most_often_small_wager = most_often_small_wager.replace(')', '')
    
    most_often_small_wager = most_often_small_wager.partition(',')[0]
    most_often_small_wager = "{:6,.2f}".format(float(most_often_small_wager))
    most_often_small_wager = str(most_often_small_wager)
    most_often_small_wager = most_often_small_wager.replace(' ', '')    
    
    p3a = p3a.replace('small_wager', (most_often_small_wager + user_currency))
    p3b = p3b.replace('small_wager', (most_often_small_wager + user_currency))
    p3c = p3c.replace('small_wager', (most_often_small_wager + user_currency))    
    
    #p3 Inserting big wager (that was calculated before the function)
    
    p3a = p3a.replace('UID', str(screen_name))
    p3b = p3b.replace('UID', str(screen_name))
    p3c = p3c.replace('UID', str(screen_name))
    
    big_wager = str(big_wager)
    big_wager = big_wager.replace(' ', '')       
    
    p3a = p3a.replace('big_wager', big_wager + user_currency)
    p3b = p3b.replace('big_wager', big_wager + user_currency)
    p3c = p3c.replace('big_wager', big_wager + user_currency)    
    
    p3_random = random.choice([p3a, p3b, p3c])
    p3_paragraph = cell110.paragraphs[2]
    p3_paragraph.text = p3_random
    
    #p4 (reasoning of investigation)
    
    p4_paragraph = cell110.paragraphs[3]
    p4_paragraph.text = p4a
    
    #p5 dropdown conclusion
    
    p5_paragraph = cell110.paragraphs[4]  
    
    if big_small_var.get() == 'Small':
        p5_paragraph.text = p52a
    else:
        p5_paragraph.text = p51a
    
    #outro
    
    conclusion = conclusion.replace('UID', str(screen_name))
    
    conclusion_run = cell210.paragraphs[0].add_run(conclusion)
    cell210.add_paragraph = conclusion_run
    
    document.save('Risk Assessment Report ' + str(uid) + '.docx')
    
def Bureaucracy(event):
    #THIS PART CREATES TXT WITH EMAIL FOR KAM
    
    today = datetime.today()
    yesterday = today - timedelta(1)
    yesterday_formated = yesterday.strftime("%d.%m.20%y")  

    operator_name = worksheet.cell(6,4).value 
    ticket_number = ticket_entry.get()  
    
    email = '''Fraud Player / screen_name / brand / operator

Hello,
    
Please be informed that the new card counter was detected. Risk Assessment Report was forwarded to the licensee. More details in the ticket below:
    
https://ezugi.worketc.com/Work?EntryID=ticketnr
    
Kind Regards,
my_name'''
    
    email = email.replace('screen_name', str(screen_name))
    email = email.replace('brand', str(operator_id_digits) + ' ' + str(operatorid))
    email = email.replace('operator', str(operator_name))
    email = email.replace('ticketnr', ticket_number)
    
    if os.environ['USERNAME'] == 'Aleksandrs':
        risk_assessor = 'Aleksandrs Karsonis'
    elif os.environ['USERNAME'] == 'Arturs':
        risk_assessor = 'Arturs Lusis'
    elif os.environ['USERNAME'] == 'Alina':
        risk_assessor = 'Alina Heifeca'
    
    email = email.replace('my_name', risk_assessor)

    with open('KAM_Email_' + screen_name + '.txt', 'w') as f:
        f.write(email)
        
    #THIS PART WRITES INFO INTO LICENSEE XL
    
    licensee_workbook = xlrd.open_workbook('C:/Shared folder/FaR Shared Folder/General/Licensee risk assessments.xlsx')
    licensee_worksheet = licensee_workbook.sheet_by_index(1)

    licensee_row = int(max(licensee_worksheet.col_values(9))) + 1
    
    licensee_wb = openpyxl.load_workbook("C:/Shared folder/FaR Shared Folder/General/Licensee risk assessments.xlsx")
    licensee_wb.active = 1
    licensee_ws = licensee_wb.active
    
    licensee_ws.cell(row=licensee_row, column=1).value = str(operator_id_digits) + ' ' + str(operatorid)
    licensee_ws.cell(row=licensee_row, column=2).value = screen_name
    licensee_ws.cell(row=licensee_row, column=3).value = uid
    licensee_ws.cell(row=licensee_row, column=4).value = risk_assessor
    licensee_ws.cell(row=licensee_row, column=5).value = "Card Counting"
    licensee_ws.cell(row=licensee_row, column=6).value = datetime.now().strftime('%d.%m.%Y %H:%M:%S')
    licensee_ws.cell(row=licensee_row, column=7).value = "Risk assessment report"
    licensee_ws.cell(row=licensee_row, column=8).value = int(ticket_number)
    licensee_ws.cell(row=licensee_row, column=10).value = int(licensee_row)
    
    licensee_wb.save("C:/Shared folder/FaR Shared Folder/General/Licensee risk assessments.xlsx")
    
    #THIS PART WRITES INFO INTO FRAUD PLAYERS XL
    
    fraud_workbook = xlrd.open_workbook('C:/Shared folder/FaR Shared Folder/General/Fraud players.xlsx')
    fraud_worksheet = fraud_workbook.sheet_by_index(0)

    fraud_row = int(max(fraud_worksheet.col_values(7))) + 1
    
    fraud_wb = openpyxl.load_workbook("C:/Shared folder/FaR Shared Folder/General/Fraud players.xlsx")
    fraud_wb.active = 0
    fraud_ws = fraud_wb.active
    
    fraud_ws.cell(row=fraud_row, column=1).value = time.strftime("%d.%m.20%y")
    fraud_ws.cell(row=fraud_row, column=2).value = str(operator_id_digits) + ' ' + str(operatorid)
    fraud_ws.cell(row=fraud_row, column=3).value = screen_name
    fraud_ws.cell(row=fraud_row, column=4).value = uid
    fraud_ws.cell(row=fraud_row, column=5).value = "Card Counting"
    fraud_ws.cell(row=fraud_row, column=6).value = int(ticket_number)
    fraud_ws.cell(row=fraud_row, column=7).value = "Found in tableau for " + yesterday_formated
    fraud_ws.cell(row=fraud_row, column=8).value = int(fraud_row)
    
    fraud_wb.save("C:/Shared folder/FaR Shared Folder/General/Fraud players.xlsx")    

#TKINTER ARCHITECTURE
#Date

frame6 = Frame(root)
frame6.pack()

dates_label = Label(frame6, text = 'Timeframe of analysis (ALL HISTORY by default)')
dates_label.pack(padx = 11, pady = 7, side = LEFT)

dates_entry = Entry(frame6)
dates_entry.pack(padx = 1, pady = 7, side = RIGHT)

#Field to put link from BO with player name and Brand

frame8 = Frame(root)
frame8.pack()

link_label = Label(frame8, text = 'BO link')
link_label.pack(padx = 120, pady = 7, side = LEFT)

link_entry = Entry(frame8)
link_entry.pack(padx = 108, pady = 7, side = RIGHT)

#Big/Small paragraph Choice

frame1 = Frame(root)
frame1.pack()

big_small_label = Label(frame1, text = 'Side bet preference (BIG by default)')
big_small_label.pack(padx = 46, pady = 7, side = LEFT)

big_small_list = ['Big', 'Small']
big_small_var = StringVar(frame1)
big_small_var.set('Click to choose')

big_small_optionmenu = OptionMenu(frame1, big_small_var, *big_small_list)
big_small_optionmenu.pack(padx = 33, side = RIGHT)

#Ticket ID

frame9 = Frame(root)
frame9.pack()

ticket_label = Label(frame9, text = 'Ticket number')
ticket_label.pack(padx = 101, pady = 7, side = LEFT)

ticket_entry = Entry(frame9)
ticket_entry.pack(padx = 88, pady = 7, side = RIGHT)

#Generate

frame_last = Frame(root)
frame_last.pack()
generate = Button(frame_last, text = 'Generate')
generate.pack(fill = X, padx = 130, pady = 10)
generate.bind('<Button-1>', Ezugi_function)

frame_10 = Frame(root)
frame_10.pack()
done = Button(frame_10, text = 'Press when report is sent')
done.pack(fill = X, padx = 130, pady = 10)
done.bind('<Button-1>', Bureaucracy)

#Goose

goose_button = Button(root, text ='Do not click here', command = rapping_goose).pack(padx = 10, side = LEFT)

root.mainloop()